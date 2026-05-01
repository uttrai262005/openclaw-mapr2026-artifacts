import json
import os
import re
from pathlib import Path
from typing import Any, Dict, List, Tuple

import urllib.request
from openpyxl import load_workbook

WS = Path(__file__).resolve().parent
BT4_DIR = WS / "dataset_clean" / "BT4"
RUBRIC_PATH = WS / "rubric" / "rubric_BT4.docx"
OUT_PATH = WS / "output" / "ket_qua_BT4_full.xlsx"
OUT_PATH_RESUME = WS / "output" / "ket_qua_BT4_full_resume.xlsx"

MAX_TC1 = 3.0
MAX_TC2 = 4.0
MAX_TC3 = 2.5
MAX_TC4 = 0.5
MAX_TOTAL = 10.0

MSSV_RE = re.compile(r"(\d{6,})")


def round_quarter(x: float) -> float:
    return round(x * 4) / 4


def round_half(x: float) -> float:
    return round(x * 2) / 2


def clamp(x: float, lo: float, hi: float) -> float:
    return max(lo, min(hi, x))


def extract_mssv(path: Path) -> str:
    m = MSSV_RE.search(path.stem)
    return m.group(1) if m else path.stem


def read_docx_text(path: Path) -> str:
    from docx import Document

    d = Document(str(path))
    parts: List[str] = []
    for p in d.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for tb in d.tables:
        for row in tb.rows:
            cells = [(c.text or "").strip().replace("\n", " ") for c in row.cells]
            line = " | ".join([c for c in cells if c])
            if line:
                parts.append(line)
    return "\n".join(parts).strip()


def read_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(str(path))
        texts: List[str] = []
        for page in reader.pages:
            t = (page.extract_text() or "").strip()
            if t:
                texts.append(t)
        return "\n\n".join(texts).strip()
    except Exception:
        pass

    from PyPDF2 import PdfReader  # type: ignore

    reader = PdfReader(str(path))
    texts = []
    for page in reader.pages:
        t = (page.extract_text() or "").strip()
        if t:
            texts.append(t)
    return "\n\n".join(texts).strip()


def read_submission_text(path: Path) -> str:
    suf = path.suffix.lower()
    if suf == ".docx":
        return read_docx_text(path)
    if suf == ".pdf":
        return read_pdf_text(path)
    return path.read_text(encoding="utf-8", errors="ignore").strip()


def get_gateway_config() -> Tuple[str, str, str]:
    url = os.environ.get("OPENCLAW_GATEWAY_URL", "").strip()
    token = os.environ.get("OPENCLAW_GATEWAY_TOKEN", "").strip()
    model = os.environ.get("OPENCLAW_MODEL", "").strip() or "openai-codex/gpt-5.2"

    cfg_path = Path(os.path.expanduser("~")) / ".openclaw" / "openclaw.json"
    cfg: Dict[str, Any] = {}
    if cfg_path.exists():
        try:
            cfg = json.loads(cfg_path.read_text(encoding="utf-8"))
        except Exception:
            cfg = {}

    def deep_get(d: Any, keys: List[str]) -> Any:
        cur = d
        for k in keys:
            if not isinstance(cur, dict) or k not in cur:
                return None
            cur = cur[k]
        return cur

    if not url:
        port = deep_get(cfg, ["gateway", "port"]) or deep_get(cfg, ["gateway", "http", "port"]) or 8317
        url = f"http://127.0.0.1:{int(port)}/v1/chat/completions"
    else:
        if url.endswith("/v1"):
            url = url + "/chat/completions"
        elif url.endswith("/v1/"):
            url = url[:-1] + "/chat/completions"
        elif url.endswith("/chat/completions"):
            pass
        else:
            url = url.rstrip("/") + "/v1/chat/completions"

    if not token:
        token = deep_get(cfg, ["gateway", "auth", "token"]) or deep_get(cfg, ["gateway", "token"]) or ""
        token = str(token) if token is not None else ""

    if not token:
        raise RuntimeError("Missing gateway token. Set OPENCLAW_GATEWAY_TOKEN or configure gateway.auth.token")

    return url, token, model


def call_llm_once(url: str, token: str, model: str, rubric_text: str, submission_text: str) -> Dict[str, Any]:
    system = (
        "Bạn là giám khảo chấm bài BT4. Chỉ dùng đúng rubric được cung cấp. "
        "Chấm đủ 4 tiêu chí trong MỘT lần. "
        "Trả về DUY NHẤT một JSON hợp lệ (không markdown, không giải thích ngoài JSON)."
    )
    user = (
        "RUBRIC (trích từ rubric_BT4.docx):\n"
        f"{rubric_text}\n\n"
        "BÀI LÀM CỦA SINH VIÊN (text trích xuất):\n"
        f"{submission_text}\n\n"
        "YÊU CẦU OUTPUT JSON với schema:\n"
        "{\n"
        "  \"tc1\": <number 0..3>,\n"
        "  \"tc2\": <number 0..4>,\n"
        "  \"tc3\": <number 0..2.5>,\n"
        "  \"tc4\": <number 0..0.5>,\n"
        "  \"nhan_xet_ngan\": <string ngắn 1-2 câu, tiếng Việt>\n"
        "}\n"
        "Chỉ xuất JSON."
    )

    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "temperature": 0.1,
        "max_tokens": 800,
    }

    req = urllib.request.Request(
        url,
        data=json.dumps(payload).encode("utf-8"),
        headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {token}",
        },
        method="POST",
    )

    with urllib.request.urlopen(req, timeout=180) as resp:
        raw = resp.read().decode("utf-8", errors="ignore")

    data = json.loads(raw)
    content = data["choices"][0]["message"]["content"]
    if not isinstance(content, str):
        content = str(content)
    content = content.strip()

    try:
        return json.loads(content)
    except Exception:
        m = re.search(r"\{[\s\S]*\}", content)
        if not m:
            raise RuntimeError("Model did not return JSON")
        return json.loads(m.group(0))


def load_rubric_text() -> str:
    text = read_docx_text(RUBRIC_PATH)
    if not text:
        raise RuntimeError("Rubric extracted empty")
    return text[:12000]


def build_index(ws) -> Dict[str, int]:
    # returns MSSV -> row index (1-based)
    idx: Dict[str, int] = {}
    for r in range(2, ws.max_row + 1):
        v = ws.cell(r, 1).value
        if v is None:
            continue
        mssv = str(v).strip()
        if mssv and mssv not in idx:
            idx[mssv] = r
    return idx


def row_is_error(ws, r: int) -> bool:
    note = ws.cell(r, 8).value
    total = ws.cell(r, 6).value
    if note and str(note).strip().lower().startswith("lỗi"):
        return True
    if total is None or str(total).strip() == "":
        return True
    return False


def main(limit: int = 30, out_path: Path = OUT_PATH):
    if not out_path.exists():
        raise SystemExit(f"Missing output: {out_path}")

    wb = load_workbook(out_path)
    ws = wb.active
    idx = build_index(ws)

    files = sorted([p for p in BT4_DIR.rglob('*') if p.is_file() and p.suffix.lower() in {'.pdf','.docx','.txt','.md'}])
    target: List[Path] = []
    for p in files:
        mssv = extract_mssv(p)
        r = idx.get(mssv)
        if r is None:
            target.append(p)
        else:
            if row_is_error(ws, r):
                target.append(p)

    # de-dup by mssv while keeping first path
    seen = set()
    uniq: List[Path] = []
    for p in target:
        mssv = extract_mssv(p)
        if mssv in seen:
            continue
        seen.add(mssv)
        uniq.append(p)

    print(f"BT4 submissions: {len(files)}")
    print(f"Need grade/regrade: {len(uniq)}")

    url, token, model = get_gateway_config()
    rubric_text = load_rubric_text()

    done = 0
    for p in uniq:
        if done >= limit:
            break
        mssv = extract_mssv(p)
        try:
            sub_text = read_submission_text(p)
            if not sub_text or len(sub_text.strip()) < 20:
                raise RuntimeError("Empty/too-short extracted text")
        except Exception:
            # update or append as error
            r = idx.get(mssv)
            if r is None:
                ws.append([mssv, "", "", "", "", "", "", "Lỗi"])
                idx[mssv] = ws.max_row
            else:
                ws.cell(r, 2).value = ""
                ws.cell(r, 3).value = ""
                ws.cell(r, 4).value = ""
                ws.cell(r, 5).value = ""
                ws.cell(r, 6).value = ""
                ws.cell(r, 7).value = ""
                ws.cell(r, 8).value = "Lỗi"
            try:
                wb.save(out_path)
            except PermissionError:
                out_path = OUT_PATH_RESUME
                wb.save(out_path)
            print(f"{mssv} {p.name}: READ ERROR -> recorded")
            done += 1
            continue

        try:
            res = call_llm_once(url, token, model, rubric_text, sub_text)
            tc1 = round_quarter(clamp(float(res.get('tc1', 0.0)), 0.0, MAX_TC1))
            tc2 = round_quarter(clamp(float(res.get('tc2', 0.0)), 0.0, MAX_TC2))
            tc3 = round_quarter(clamp(float(res.get('tc3', 0.0)), 0.0, MAX_TC3))
            tc4 = round_quarter(clamp(float(res.get('tc4', 0.0)), 0.0, MAX_TC4))
            total = round_half(clamp(tc1 + tc2 + tc3 + tc4, 0.0, MAX_TOTAL))
            short = str(res.get('nhan_xet_ngan', '')).strip()

            r = idx.get(mssv)
            if r is None:
                ws.append([mssv, tc1, tc2, tc3, tc4, total, short, ""])
                idx[mssv] = ws.max_row
            else:
                ws.cell(r, 2).value = tc1
                ws.cell(r, 3).value = tc2
                ws.cell(r, 4).value = tc3
                ws.cell(r, 5).value = tc4
                ws.cell(r, 6).value = total
                ws.cell(r, 7).value = short
                ws.cell(r, 8).value = ""

            try:
                wb.save(out_path)
            except PermissionError:
                out_path = OUT_PATH_RESUME
                wb.save(out_path)
            print(f"{mssv} {p.name}: OK total={total}")
        except Exception as e:
            r = idx.get(mssv)
            if r is None:
                ws.append([mssv, "", "", "", "", "", "", "Lỗi"])
                idx[mssv] = ws.max_row
            else:
                ws.cell(r, 2).value = ""
                ws.cell(r, 3).value = ""
                ws.cell(r, 4).value = ""
                ws.cell(r, 5).value = ""
                ws.cell(r, 6).value = ""
                ws.cell(r, 7).value = ""
                ws.cell(r, 8).value = "Lỗi"
            try:
                wb.save(out_path)
            except PermissionError:
                out_path = OUT_PATH_RESUME
                wb.save(out_path)
            print(f"{mssv} {p.name}: API/JSON ERROR -> recorded ({type(e).__name__})")

        done += 1

    print(f"DONE this run: {done}")
    print(f"Output file: {out_path}")


if __name__ == '__main__':
    import argparse

    ap = argparse.ArgumentParser()
    ap.add_argument('--limit', type=int, default=30)
    ap.add_argument('--out', type=str, default=str(OUT_PATH), help='xlsx to update (default existing ket_qua_BT4_full.xlsx)')
    args = ap.parse_args()
    main(limit=args.limit, out_path=Path(args.out))
