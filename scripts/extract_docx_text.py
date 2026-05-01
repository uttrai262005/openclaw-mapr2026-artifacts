import sys
from pathlib import Path

try:
    sys.stdout.reconfigure(encoding='utf-8')
except Exception:
    pass

import docx

p = Path(sys.argv[1])
d = docx.Document(str(p))

def emit(text: str):
    t = (text or '').strip()
    if t:
        print(t)

# paragraphs
for para in d.paragraphs:
    emit(para.text)

# tables
for ti, table in enumerate(d.tables, start=1):
    emit(f"\n--- TABLE {ti} ---")
    for row in table.rows:
        cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
        line = ' | '.join([c for c in cells if c])
        emit(line)
