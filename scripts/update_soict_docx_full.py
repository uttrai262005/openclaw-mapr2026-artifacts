"""Update SOICT Word draft with concrete model config, rubric/assignment summaries, and stronger Results narrative.

This script edits the existing paper/SOICT_draft_ccis.docx by inserting text blocks.
(It does not yet add Related Work citations; that will be a separate step.)
"""

import pandas as pd
from pathlib import Path
from docx import Document

DOC_PATH = Path('paper/SOICT_draft_ccis.docx')
SUM_PATH = Path('output/soict_analysis_tables.xlsx')

summary = pd.read_excel(SUM_PATH, sheet_name='summary')

MODEL_BLOCK = (
    "Implementation details (reproducibility). We served all LLM calls through the OpenClaw Gateway (OpenAI-compatible /v1/chat/completions). "
    "Both Phase 1 and Phase 2 used the same underlying model: gpt-5.2 accessed via Codex OAuth within OpenClaw. "
    "In Phase 1 scripts (e.g., bt1_grade_incremental.py), the request sets temperature=0.1 and max_tokens=900. "
    "In Phase 2 graders (giang-vien-v2), the request sets temperature=0.1 and max_tokens=800, with a 180s timeout and retry logic for rate limits and JSON parsing. "
    "All graders are constrained to return JSON-only outputs with per-criterion scores, short comments, and (optionally) bullet highlights/issues."
)

ASSIGNMENT_BLOCK = (
    "Assignments and rubrics. The dataset consists of four rubric-based assignments in an undergraduate e-commerce introduction course at UIT. "
    "BT1 asks students to study the E-commerce program curriculum and learning outcomes; its rubric contains 5 criterion groups (2.0 points each) emphasizing factual correctness and sourcing. "
    "BT2 asks students to survey the e-commerce job market for two positions and analyze job descriptions, skills, and gaps; its rubric contains 5 criteria with maxima 1.5/2.0/2.5/2.0/2.0. "
    "BT3 requires two CVs tailored to two positions; its rubric contains 5 criteria with maxima 1.5/2.0/3.0/2.0/1.5. "
    "BT4 is an essay on personal career roadmap with required structure; its rubric contains 4 criteria with maxima 3.0/4.0/2.5/0.5."
)

BT12_EXPLAIN = (
    "Why BT1–BT2 are harder. Compared to BT3–BT4, BT1–BT2 emphasize factual completeness, correct references, and domain-specific details (curriculum, official sources, job-market evidence). "
    "These tasks are more sensitive to missing citations, partial answers, and rubric interpretation variance, which can reduce AI-vs-human ordinal agreement (QWK). "
    "In addition, Phase 2 decomposition can overweight surface features (writing mechanics) relative to rubric intent in BT1–BT2, contributing to reduced QWK."
)

# Load doc
if not DOC_PATH.exists():
    raise SystemExit(f'Missing {DOC_PATH}')

doc = Document(DOC_PATH)

# Helper: find paragraph containing text

def find_paragraph_index(needle: str):
    for i, p in enumerate(doc.paragraphs):
        if needle in p.text:
            return i
    return -1

# Insert model block after heading '3.2 Phase 1: Single-agent rubric grading'
idx = find_paragraph_index('3.2 Phase 1: Single-agent rubric grading')
if idx != -1:
    # insert after the next paragraph (the TO DO line)
    insert_at = idx + 2
    doc.paragraphs[insert_at-1].insert_paragraph_before(MODEL_BLOCK)

# Insert assignment block after '4 Experimental Setup'
idx2 = find_paragraph_index('4 Experimental Setup')
if idx2 != -1:
    insert_at = idx2 + 1
    doc.paragraphs[insert_at].insert_paragraph_before(ASSIGNMENT_BLOCK)

# Insert BT1/BT2 explanation after '6 Discussion & Recommendations'
idx3 = find_paragraph_index('6 Discussion & Recommendations')
if idx3 != -1:
    insert_at = idx3 + 1
    doc.paragraphs[insert_at].insert_paragraph_before(BT12_EXPLAIN)

# Save updated doc
out = Path('paper/SOICT_draft_ccis_v2.docx')
doc.save(out)
print('Wrote:', out.resolve())
