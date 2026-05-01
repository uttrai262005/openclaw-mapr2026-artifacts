from pathlib import Path
import sys

try:
    sys.stdout.reconfigure(encoding='utf-8')
except Exception:
    pass

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'output' / '4_bai_mau_theo_rubric_10d.docx'


def add_title(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_h(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_score_table(doc, rows):
    # rows: list of (criterion, max, score, note)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = 'Tiêu chí'
    hdr[1].text = 'Điểm tối đa'
    hdr[2].text = 'Điểm mẫu'
    hdr[3].text = 'Ghi chú chấm'
    for crit, mx, sc, note in rows:
        cells = table.add_row().cells
        cells[0].text = str(crit)
        cells[1].text = str(mx)
        cells[2].text = str(sc)
        cells[3].text = str(note)
    doc.add_paragraph()


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    add_title(doc, '04 BÀI MẪU (10 ĐIỂM) THEO RUBRIC — EC005 (TMĐT)')
    doc.add_paragraph('Mục đích: dùng làm “anchor” để chấm tay nhất quán theo rubric 10 điểm. ')
    doc.add_paragraph('Lưu ý: Đây là bài mẫu tham khảo cách trình bày + mức độ chi tiết + dẫn nguồn. Có thể thay link nguồn bằng link chính thống tương đương.')

    # ---------------- BT1 ----------------
    add_h(doc, 'BT1 — Bài mẫu (mức 10/10)', level=1)
    doc.add_paragraph('Chủ đề: Tìm hiểu UIT + CTĐT ngành TMĐT (trả lời theo nhóm câu hỏi trong đề).')

    add_h(doc, 'Bài làm mẫu', level=2)
    doc.add_paragraph('Nguồn tham khảo chính (trích dẫn trong bài):')
    doc.add_paragraph('- UIT: https://www.uit.edu.vn (trang giới thiệu/chiến lược/phần sứ mệnh–tầm nhìn tuỳ thời điểm công bố).')
    doc.add_paragraph('- Đào tạo (DAA): https://daa.uit.edu.vn')
    doc.add_paragraph('- CTĐT/chuẩn đầu ra ngành TMĐT (file CTĐT 2024 do môn học cung cấp hoặc trang DAA).')
    doc.add_paragraph('- Báo cáo TMĐT Việt Nam: Sách trắng TMĐT (VECOM) https://vecom.vn')
    doc.add_paragraph('- Báo cáo quốc tế: eMarketer/Insider Intelligence https://www.insiderintelligence.com , Statista https://www.statista.com')

    doc.add_paragraph('1) Tầm nhìn – sứ mệnh – triết lý giáo dục của UIT (tóm tắt):')
    doc.add_paragraph('- Tầm nhìn: trở thành trường đại học hàng đầu về CNTT&TT, có vị thế trong khu vực/quốc tế; đóng góp tri thức và đổi mới sáng tạo. (Nguồn: UIT)')
    doc.add_paragraph('- Sứ mệnh: đào tạo nguồn nhân lực chất lượng cao; nghiên cứu – chuyển giao; phục vụ cộng đồng/đóng góp phát triển kinh tế số. (Nguồn: UIT)')
    doc.add_paragraph('- Triết lý/định hướng giáo dục: học đi đôi với hành; gắn đào tạo với nhu cầu doanh nghiệp; đề cao đạo đức học thuật và năng lực tự học. (Nguồn: UIT)')

    doc.add_paragraph('2) Mục tiêu đào tạo của UIT:')
    doc.add_paragraph('- Trang bị nền tảng kiến thức vững; năng lực nghề nghiệp theo chuẩn đầu ra; kỹ năng mềm và ngoại ngữ; thái độ – đạo đức nghề nghiệp. (Nguồn: UIT)')

    doc.add_paragraph('3) Mục tiêu đào tạo ngành TMĐT (tóm tắt theo CTĐT):')
    doc.add_paragraph('- Kiến thức: nền tảng CNTT + kiến thức nghiệp vụ TMĐT (marketing số, vận hành sàn, logistics, dữ liệu, pháp lý…).')
    doc.add_paragraph('- Kỹ năng: phân tích dữ liệu/khách hàng, thiết kế & triển khai chiến dịch, vận hành sản phẩm TMĐT, phối hợp liên phòng ban.')
    doc.add_paragraph('- Thái độ: tư duy khách hàng, tuân thủ pháp lý & đạo đức dữ liệu, học suốt đời. (Nguồn: CTĐT/DAA)')

    doc.add_paragraph('4) Chuẩn đầu ra ngành TMĐT (minh hoạ nhóm năng lực):')
    doc.add_paragraph('- Năng lực chuyên môn: hiểu hệ thống TMĐT, hành vi người tiêu dùng, mô hình kinh doanh số.')
    doc.add_paragraph('- Năng lực CNTT/dữ liệu: sử dụng công cụ phân tích (Excel/SQL/BI), hiểu nền tảng web/app, bảo mật cơ bản.')
    doc.add_paragraph('- Năng lực triển khai: lập kế hoạch, thực thi, đo lường KPI; làm việc nhóm; giao tiếp; tiếng Anh. (Nguồn: CTĐT/DAA)')

    doc.add_paragraph('5) Học TMĐT ra trường làm gì? (ví dụ vị trí):')
    doc.add_paragraph('- E-commerce Executive / E-commerce Specialist (vận hành gian hàng/sàn).')
    doc.add_paragraph('- Digital Marketing / Performance Marketing (chạy ads, tối ưu chuyển đổi).')
    doc.add_paragraph('- CRM/Customer Experience, Sales Ops, Merchandising, Content Commerce.')
    doc.add_paragraph('- Data Analyst (TMĐT), Product/Business Analyst (mảng tăng trưởng).')

    doc.add_paragraph('6) Tổng số tín chỉ tối thiểu toàn khóa: theo CTĐT Khoá 2024 (trích đúng con số trong file CTĐT). (Nguồn: CTĐT/DAA)')
    doc.add_paragraph('7) Điều kiện tốt nghiệp: hoàn thành đủ tín chỉ, đạt chuẩn ngoại ngữ, hoàn thành khối kiến thức tốt nghiệp/đồ án (nếu có) theo quy định. (Nguồn: CTĐT/DAA)')
    doc.add_paragraph('8) Chuẩn đầu ra tiếng Anh Khoá 2024: nêu chuẩn theo quy định CTĐT (ví dụ tương đương IELTS/TOEIC/CEFR theo văn bản). (Nguồn: CTĐT/DAA)')

    doc.add_paragraph('9) Môn/học phần nhóm CNTT (ví dụ, liệt kê theo CTĐT):')
    doc.add_paragraph('- Nhập môn lập trình; Cấu trúc dữ liệu & giải thuật; Cơ sở dữ liệu; Mạng máy tính; Kỹ nghệ phần mềm; Phân tích thiết kế HTTT; An toàn thông tin cơ bản… (Nguồn: CTĐT)')

    doc.add_paragraph('10) Môn/học phần nhóm nghiệp vụ cốt lõi TMĐT (ví dụ):')
    doc.add_paragraph('- Marketing căn bản/Marketing số; Quản trị bán hàng; Quản trị chuỗi cung ứng/logistics; Quản trị quan hệ khách hàng; Thanh toán điện tử; Thương mại điện tử; Phân tích dữ liệu trong kinh doanh… (Nguồn: CTĐT)')

    doc.add_paragraph('11) Kỹ năng mềm/đạo đức:')
    doc.add_paragraph('- Ở trường: thuyết trình, viết báo cáo, teamwork qua môn học/dự án; tham gia CLB học thuật. Ngoài trường: intern/part-time, hoạt động cộng đồng, hackathon. Đạo đức: tôn trọng bản quyền, trích dẫn nguồn, không “data scraping” trái phép, bảo vệ dữ liệu cá nhân.')

    doc.add_paragraph('12) Tự chọn tự do: học phần ngoài khung bắt buộc để mở rộng năng lực (ví dụ: ngoại ngữ, thiết kế, phân tích dữ liệu nâng cao) theo quy định đăng ký tín chỉ. (Nguồn: CTĐT/DAA)')
    doc.add_paragraph('13) Khối kiến thức tốt nghiệp (12TC): mô tả cấu phần (đồ án/khóa luận/thực tập) theo CTĐT; đầu ra mong đợi (báo cáo, demo, đánh giá). (Nguồn: CTĐT)')

    doc.add_paragraph('16) Nền tảng học trực tuyến:')
    doc.add_paragraph('- Việt Nam: Kyna, Unica (chọn lọc), FUNiX (tuỳ chương trình).')
    doc.add_paragraph('- Quốc tế: Coursera, edX, Udemy (chọn khoá chất lượng), Google Digital Garage, LinkedIn Learning.')

    doc.add_paragraph('17) Nguồn báo cáo/thống kê TMĐT uy tín:')
    doc.add_paragraph('- Việt Nam: VECOM (Sách trắng TMĐT), Bộ Công Thương (Cục TMĐT&KTS), Tổng cục Thống kê (GSO) cho chỉ số vĩ mô.')
    doc.add_paragraph('- Quốc tế: UNCTAD eCommerce, World Bank Data, OECD, Statista (tham khảo), Insider Intelligence/eMarketer.')

    doc.add_paragraph('18) Tạp chí/kỷ yếu/khoá luận nên đọc:')
    doc.add_paragraph('- Google Scholar (lọc nguồn), ACM Digital Library, IEEE Xplore (tuỳ chủ đề), SpringerLink; các proceedings về IS/marketing analytics. Ở UIT: kho luận văn/đồ án (nếu được truy cập).')

    add_h(doc, 'Bảng chấm mẫu (đạt 10/10 theo rubric)', level=2)
    add_score_table(doc, [
        ('Nhóm 1 (C1–2): UIT tầm nhìn/sứ mệnh/triết lý + mục tiêu', 2.0, 2.0, 'Đủ ý + diễn giải + có nguồn UIT.'),
        ('Nhóm 2 (C3–4): Ngành TMĐT mục tiêu & chuẩn đầu ra', 2.0, 2.0, 'Tóm tắt đúng, có nguồn CTĐT/DAA.'),
        ('Nhóm 3 (C5–8): Việc làm + tín chỉ + tốt nghiệp + chuẩn TA', 2.0, 2.0, 'Trả lời đủ 4 câu, có ví dụ nghề.'),
        ('Nhóm 4 (C9–10): Liệt kê học phần CNTT & nghiệp vụ', 2.0, 2.0, 'Phân nhóm rõ, bám CTĐT.'),
        ('Nhóm 5 (C11–13,16–18): Kỹ năng mềm + tự chọn + 12TC + nguồn học/báo cáo/tạp chí', 2.0, 2.0, 'Ví dụ cụ thể + có nguồn uy tín.'),
    ])

    # ---------------- BT2 ----------------
    add_h(doc, 'BT2 — Bài mẫu (mức 10/10)', level=1)
    doc.add_paragraph('Chủ đề: Khảo sát thị trường việc làm ngành TMĐT theo hướng dẫn repo.')

    add_h(doc, 'Bài làm mẫu', level=2)
    doc.add_paragraph('2 vị trí lựa chọn:')
    doc.add_paragraph('A) E-commerce Operations Specialist (Vận hành sàn/Gian hàng).')
    doc.add_paragraph('B) Performance Marketing Executive (Quảng cáo chuyển đổi).')
    doc.add_paragraph('Lý do phù hợp: đúng “core” TMĐT (vận hành kênh bán + tăng trưởng), tận dụng kiến thức CTĐT: marketing số, dữ liệu, vận hành, chăm sóc khách hàng.')

    doc.add_paragraph('Thu thập JD (mỗi vị trí ≥2 nguồn) — ví dụ:')
    doc.add_paragraph('- Vị trí A:')
    doc.add_paragraph('  (1) TopCV: https://www.topcv.vn/ (tìm “E-commerce Executive/Operation”) — truy cập 29/03/2026')
    doc.add_paragraph('  (2) VietnamWorks: https://www.vietnamworks.com/ (tìm “Ecommerce Operation”) — truy cập 29/03/2026')
    doc.add_paragraph('- Vị trí B:')
    doc.add_paragraph('  (1) LinkedIn Jobs: https://www.linkedin.com/jobs/ (tìm “Performance Marketing”) — truy cập 29/03/2026')
    doc.add_paragraph('  (2) YBOX/BrandsVietnam Jobs: https://jobs.brandsvietnam.com/ — truy cập 29/03/2026')

    doc.add_paragraph('Phân tích yêu cầu tuyển dụng (tổng hợp từ JD):')
    doc.add_paragraph('1) Hard skills:')
    doc.add_paragraph('- A (Ops): quản trị sản phẩm trên sàn, merchandising, pricing/promo, tồn kho, xử lý đơn, CSKH, phối hợp fulfillment.')
    doc.add_paragraph('- B (Perf): chạy ads (Meta/Google/TikTok), tối ưu funnel, A/B testing, đo lường ROAS/CAC, remarketing.')
    doc.add_paragraph('2) Tools:')
    doc.add_paragraph('- Excel/Google Sheets (pivot, vlookup), dashboard; GA4; nền tảng ads; công cụ quản lý sàn (Seller Center/Shopee/Lazada… tuỳ JD).')
    doc.add_paragraph('3) Soft skills: tư duy số liệu, giao tiếp liên phòng ban, chịu áp lực KPI, chủ động học nhanh.')
    doc.add_paragraph('4) Điểm chung vs khác:')
    doc.add_paragraph('- Chung: KPI-driven, phân tích dữ liệu, hiểu hành vi mua sắm, giao tiếp tốt.')
    doc.add_paragraph('- Khác: Ops thiên về vận hành – quy trình – phối hợp kho/CS; Perf thiên về media buying – creative – tối ưu chuyển đổi.')

    doc.add_paragraph('Đối chiếu với CTĐT UIT & khoảng trống kỹ năng (ví dụ trình bày):')
    doc.add_paragraph('- Ở UIT: nền tảng marketing/kinh doanh số, phân tích dữ liệu, dự án môn học giúp luyện viết báo cáo, thuyết trình.')
    doc.add_paragraph('- Gap cần bù:')
    doc.add_paragraph('  + GA4/Tracking thực chiến;')
    doc.add_paragraph('  + Kỹ năng chạy ads và tối ưu ngân sách;')
    doc.add_paragraph('  + Portfolio case study (campaign report / shop audit).')
    doc.add_paragraph('- Kế hoạch bù (cụ thể): 6–8 tuần học Google Skillshop + làm 2 mini-project: (1) audit gian hàng giả lập, (2) chạy thử campaign mô phỏng + báo cáo KPI.')

    doc.add_paragraph('Bổ sung khảo sát thêm (có nguồn):')
    doc.add_paragraph('- Dải lương tham khảo theo thị trường (nguồn: báo cáo lương TopCV/VietnamWorks hoặc JobStreet).')
    doc.add_paragraph('- Chứng chỉ: Google Ads/GA4, Meta Blueprint (Perf); Shopee/Lazada seller training (Ops).')
    doc.add_paragraph('- Lộ trình: Junior → Senior → Team lead (theo competency & KPI).')

    add_h(doc, 'Bảng chấm mẫu (đạt 10/10 theo rubric)', level=2)
    add_score_table(doc, [
        ('1. Chọn 02 vị trí phù hợp', 1.5, 1.5, '2 vị trí rõ + giải thích hợp lý.'),
        ('2. Thu thập JD/nguồn', 2.0, 2.0, 'Mỗi vị trí ≥2 nguồn + có link/ngày truy cập.'),
        ('3. Phân tích yêu cầu theo hướng dẫn', 2.5, 2.5, 'Có cấu trúc + tổng hợp chung/khác.'),
        ('4. Đối chiếu CTĐT & gap', 2.0, 2.0, 'Nêu gap + kế hoạch bù đo được.'),
        ('5. Bổ sung khảo sát thêm', 2.0, 2.0, 'Có dữ liệu/chứng chỉ/xu hướng + kết luận.'),
    ])

    # ---------------- BT3 ----------------
    add_h(doc, 'BT3 — Bài mẫu (mức 10/10)', level=1)
    doc.add_paragraph('Chủ đề: 02 CV cho 02 vị trí (dùng làm chuẩn chấm theo rubric).')

    add_h(doc, 'CV Mẫu 1 — E-commerce Operations Specialist', level=2)
    doc.add_paragraph('Họ tên: Nguyễn Văn A | Email: a.nguyen@email.com | SĐT: 09xx | LinkedIn: linkedin.com/in/nguyenvana | Portfolio: notion.site/nguyenvana')
    doc.add_paragraph('Mục tiêu: Ứng tuyển vị trí E-commerce Operations Specialist. Mục tiêu 6 tháng: nắm vận hành Shopee/Lazada, cải thiện CR +5% và giảm hủy đơn -10% thông qua tối ưu listing & quy trình xử lý đơn.')
    doc.add_paragraph('Kỹ năng chính: Excel (Pivot, PowerQuery cơ bản), phân tích KPI sàn (CR, AOV, refund), merchandising, CSKH, phối hợp fulfillment.')
    doc.add_paragraph('Dự án (minh hoạ impact):')
    doc.add_paragraph('- Shop Audit Case Study: phân tích 200 SKU (giả lập) → đề xuất tối ưu tiêu đề/ảnh/thuộc tính; mô phỏng tăng CTR +12% (từ benchmark).')
    doc.add_paragraph('- Dashboard vận hành: xây bảng theo dõi đơn hàng, SLA, hoàn hủy; cảnh báo SKU tồn kho thấp bằng conditional formatting.')
    doc.add_paragraph('Hoạt động: CLB/đội nhóm; vai trò điều phối; kỹ năng teamwork.')

    add_h(doc, 'CV Mẫu 2 — Performance Marketing Executive', level=2)
    doc.add_paragraph('Họ tên: Nguyễn Văn A | Email/SĐT/LinkedIn/Portfolio như trên')
    doc.add_paragraph('Mục tiêu: Ứng tuyển Performance Marketing Executive. Mục tiêu 3 tháng: chạy & tối ưu chiến dịch chuyển đổi cho sản phẩm TMĐT, đạt ROAS mục tiêu và xây 02 case study đo lường rõ ràng.')
    doc.add_paragraph('Kỹ năng chính: Google Ads/Meta Ads (nền tảng), GA4 (report), đọc insight, A/B testing creative, tối ưu landing page cơ bản.')
    doc.add_paragraph('Dự án (minh hoạ impact):')
    doc.add_paragraph('- Campaign mô phỏng: thiết kế funnel (Awareness→Conversion), set KPI (CTR, CPA, ROAS); báo cáo tuần, đề xuất tối ưu audience/creative.')
    doc.add_paragraph('- Phân tích dữ liệu: dùng Excel/SQL cơ bản để phân khúc khách hàng RFM (giả lập) → đề xuất remarketing.')

    add_h(doc, 'Bảng chấm mẫu (đạt 10/10 theo rubric)', level=2)
    add_score_table(doc, [
        ('1. Có 02 CV cho 02 vị trí', 1.5, 1.5, '2 CV khác mục tiêu + kỹ năng trọng tâm.'),
        ('2. Nội dung cơ bản đầy đủ', 2.0, 2.0, 'Có thông tin, học vấn/kỹ năng/dự án/liên hệ.'),
        ('3. Mapping CV ↔ JD', 3.0, 3.0, 'Nêu kỹ năng/công cụ/impact bám JD.'),
        ('4. Trình bày/format', 2.0, 2.0, 'Rõ ràng, có link portfolio/LinkedIn.'),
        ('5. Ngôn ngữ & thông điệp', 1.5, 1.5, 'Mục tiêu cụ thể, tránh sáo rỗng.'),
    ])

    # ---------------- BT4 ----------------
    add_h(doc, 'BT4 — Bài mẫu (mức 10/10)', level=1)
    doc.add_paragraph('Chủ đề: Tiểu luận “Lộ trình nghề nghiệp” (bố cục 1–2–3 theo yêu cầu).')

    add_h(doc, 'Bài làm mẫu', level=2)
    doc.add_paragraph('1) MỤC TIÊU NGHỀ NGHIỆP')
    doc.add_paragraph('- Vị trí mong muốn sau khi ra trường: Junior Performance Marketer (mảng TMĐT) hoặc Growth Analyst (entry-level).')
    doc.add_paragraph('- Công ty mong muốn: doanh nghiệp TMĐT/retail có data-driven (ví dụ: Tiki/Shopee Mall brands/Thế Giới Di Động mảng online…) hoặc agency performance uy tín.')
    doc.add_paragraph('- Ngôn ngữ sử dụng: tiếng Việt + tiếng Anh trong báo cáo/trao đổi công việc.')
    doc.add_paragraph('- Quốc gia mong muốn (tuỳ cơ hội): Singapore (môi trường TMĐT mạnh, gần VN) trong 3–5 năm.')
    doc.add_paragraph('- Mức lương mục tiêu: >1000$ sau 2–3 năm (khi lên Senior hoặc có portfolio chứng minh ROI).')
    doc.add_paragraph('- Học lên cao/chứng chỉ: hoàn thành chứng chỉ GA4/Google Ads/Meta; cân nhắc học ThS (Data/Marketing Analytics) sau 3–5 năm nếu phù hợp.')
    doc.add_paragraph('- Lý do: phù hợp sở thích phân tích số liệu + tối ưu tăng trưởng; thị trường TMĐT tiếp tục mở rộng; công việc có lộ trình rõ theo KPI/impact.')

    doc.add_paragraph('2) QUÁ TRÌNH CHUẨN BỊ (KSA)')
    doc.add_paragraph('2.1 Kiến thức:')
    doc.add_paragraph('- Học tại UIT: nền tảng TMĐT, marketing, dữ liệu/cơ sở dữ liệu, kỹ năng báo cáo.')
    doc.add_paragraph('- Bổ sung ngoài: tracking/attribution, hành vi khách hàng nâng cao, thống kê ứng dụng, kiến thức nền tảng về creative & copywriting.')
    doc.add_paragraph('2.2 Kỹ năng:')
    doc.add_paragraph('- Công cụ: Excel nâng cao, GA4, Looker Studio/Power BI, nền tảng ads.')
    doc.add_paragraph('- Thực chiến: đọc insight, đặt giả thuyết, A/B testing, viết báo cáo weekly/monthly.')
    doc.add_paragraph('2.3 Thái độ:')
    doc.add_paragraph('- Kỷ luật theo KPI, trung thực số liệu, tôn trọng dữ liệu cá nhân, chủ động học và nhận feedback.')
    doc.add_paragraph('2.4 Mục tiêu NĂM 1 (SMART):')
    doc.add_paragraph('- Hoàn thành 01 chứng chỉ GA4 + 01 chứng chỉ Google Ads trong 12 tuần;')
    doc.add_paragraph('- Làm 02 mini case study (shop audit + campaign mô phỏng) đăng lên portfolio;')
    doc.add_paragraph('- GPA mục tiêu ≥ 7.5, tham gia 01 CLB học thuật và thuyết trình ít nhất 02 lần.')
    doc.add_paragraph('2.5 Mục tiêu NĂM 2 (SMART):')
    doc.add_paragraph('- Thực tập 2–3 tháng ở vị trí liên quan;')
    doc.add_paragraph('- Làm 01 dự án dữ liệu: dashboard KPI TMĐT từ dataset giả lập;')
    doc.add_paragraph('- Nâng tiếng Anh: đạt tương đương B2/IELTS 6.0 (hoặc chuẩn theo CTĐT).')

    doc.add_paragraph('3) LỘ TRÌNH 2/5/10 NĂM')
    doc.add_paragraph('- 2 năm: hoàn thành nền tảng + 1 lần thực tập; có portfolio tối thiểu 4 case study; có khả năng tự chạy & báo cáo campaign nhỏ.')
    doc.add_paragraph('- 5 năm: lên Senior (hoặc Specialist) quản lý ngân sách/đầu mục tăng trưởng; dẫn dắt 1–2 bạn junior; thu nhập đạt mục tiêu; có cơ hội làm việc môi trường quốc tế (SG) nếu phù hợp.')
    doc.add_paragraph('- 10 năm: Growth Lead/Head of Performance (hoặc chuyển hướng Product Growth); xây hệ thống đo lường chuẩn; mentor đội nhóm; cân nhắc học cao học/chứng chỉ chuyên sâu (Marketing Analytics/Data).')
    doc.add_paragraph('Rủi ro & phương án dự phòng:')
    doc.add_paragraph('- Nếu thị trường ads biến động/chi phí tăng: dịch chuyển sang CRM/Retention hoặc Analytics để tối ưu LTV;')
    doc.add_paragraph('- Nếu thiếu cơ hội thực tập: làm dự án freelance/phi lợi nhuận (chạy chiến dịch nhỏ) để có dữ liệu thật;')
    doc.add_paragraph('- Nếu tiếng Anh chưa đạt: ưu tiên lịch học 30 phút/ngày + speaking group 2 buổi/tuần.')

    add_h(doc, 'Bảng chấm mẫu (đạt 10/10 theo rubric)', level=2)
    add_score_table(doc, [
        ('1. Mục tiêu nghề nghiệp (đủ gạch đầu dòng + lý do)', 3.0, 3.0, 'Cụ thể, thực tế, có lý do.'),
        ('2. Quá trình chuẩn bị (KSA + UIT vs ngoài + mục tiêu năm 1-2 SMART)', 4.0, 4.0, 'Phân loại rõ, mục tiêu đo được.'),
        ('3. Lộ trình 2/5/10 năm + milestone + backup plan', 2.5, 2.5, 'Timeline rõ, có rủi ro/backup.'),
        ('4. Trình bày & lập luận', 0.5, 0.5, 'Mạch lạc, đúng bố cục 1-2-3.'),
    ])

    doc.add_page_break()
    doc.add_paragraph('Gợi ý sử dụng khi chấm tay:')
    doc.add_paragraph('- Chấm theo từng tiêu chí, tránh “cảm tính tổng thể”.')
    doc.add_paragraph('- Với bài thật: nếu thiếu nguồn/link hoặc nguồn kém uy tín, trừ 0.25–0.5 ở tiêu chí liên quan (đúng như gợi ý rubric BT1/BT2).')

    doc.save(str(OUT))
    print(str(OUT))


if __name__ == '__main__':
    main()
