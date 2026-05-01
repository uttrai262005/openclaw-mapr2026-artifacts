import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

summary = pd.read_excel('output/soict_analysis_tables.xlsx', sheet_name='summary')

out_path = Path('paper/SOICT_draft_ccis.docx')
out_path.parent.mkdir(parents=True, exist_ok=True)

doc = Document()

# Title
p = doc.add_paragraph('An Auditable Agentic LLM Pipeline for Rubric-Based Grading: Evidence from a Real E-commerce Course Dataset')
if p.runs:
    p.runs[0].bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Trúc [Author list TBD], University of Information Technology (UIT), VNU-HCM')

doc.add_heading('Abstract', level=1)
doc.add_paragraph(
    'Large enrollment courses rely on rubric-based assessment to ensure transparency and consistency, but grading is labor-intensive and often delays feedback. '
    'We present an auditable agentic grading pipeline implemented with the OpenClaw framework and apply it to four rubric-based assignments from an undergraduate '
    'e-commerce introduction course at UIT. The pipeline includes dataset normalization, optional OCR for scanned documents, incremental result writing, and '
    'automated auditing to guarantee one-to-one mapping between submissions and exported scores. Two human graders independently score a subset of submissions; '
    'ground truth is defined as the mean of the two scores. We compare a single-agent configuration (Phase 1) with a multi-agent configuration (Phase 2) that '
    'decomposes grading into content/structure/language graders and aggregates scores via a chairman agent. Results show that multi-agent grading does not '
    'automatically improve reliability: Phase 2 can induce assignment-dependent score shifts and higher disagreement among specialized graders. We provide '
    'practical recommendations for calibrating and aggregating multi-agent rubric grading.'
)

doc.add_paragraph('Keywords: automated grading; rubric-based assessment; agentic AI; LLM; OpenClaw; reliability; quadratic weighted kappa')

# Introduction

doc.add_heading('1 Introduction', level=1)
doc.add_paragraph('We study rubric-grounded automated grading in a real classroom setting with heterogeneous submissions and strong auditability requirements.')

doc.add_heading('1.1 Motivation', level=2)
doc.add_paragraph(
    'Rubric-based grading supports fairness and consistency but is time-consuming in large classes. Submissions appear in multiple formats '
    '(DOCX/PDF/TXT; scanned artifacts), and instructors require deterministic rounding rules and complete exports for administration and research.'
)

doc.add_heading('1.2 Gap', level=2)
doc.add_paragraph(
    'Prior LLM-based grading studies often lack end-to-end auditable pipelines and do not contextualize system performance against human inter-rater agreement. '
    'It remains unclear when multi-agent decomposition helps reliability under rubric scoring in realistic course settings.'
)

doc.add_heading('1.3 Contributions', level=2)
doc.add_paragraph('C1: Auditable OpenClaw pipeline (normalize → OCR → grading → incremental export → audit).')
doc.add_paragraph('C2: Clean dataset artifacts for four assignments (UIT e-commerce intro course).')
doc.add_paragraph('C3: Reliability-first evaluation with human inter-rater QWK and AI-vs-GT (QWK/MAE/Pearson).')
doc.add_paragraph('C4: Evidence-based analysis of why multi-agent can underperform single-agent, plus practical recommendations.')

# Related work placeholder

doc.add_heading('2 Related Work (TO BE COMPLETED)', level=1)
doc.add_paragraph(
    'TO DO (blocker before submission): write 2–3 pages with ~15–20 citations covering: '
    '(i) traditional AES/SAS datasets/systems (e.g., ASAP), '
    '(ii) recent LLM-based grading/rubric scoring (2023–2025), '
    '(iii) multi-agent LLM frameworks and aggregation, and '
    '(iv) educational measurement reliability and QWK.'
)

# Method

doc.add_heading('3 Method', level=1)

doc.add_heading('3.1 Pipeline and auditability', level=2)
doc.add_paragraph(
    'Inputs: student submissions (DOCX/PDF/TXT; scanned artifacts) and rubric documents. Steps: normalization, text extraction, OCR when needed, '
    'LLM grading, deterministic rounding (0.25 per criterion; 0.5 total), incremental XLSX export, and automated audits (missingness/duplicates/score constraints).'
)

doc.add_heading('3.2 Phase 1: Single-agent rubric grading', level=2)
doc.add_paragraph(
    'TO DO (reproducibility blocker): specify the exact model (provider + version), temperature/top_p, max tokens, and prompt template structure '
    '(rubric injection + required output schema).'
)

doc.add_heading('3.3 Phase 2: Multi-agent grading (OpenClaw giang-vien-v2)', level=2)
doc.add_paragraph(
    'Phase 2 runs three graders in parallel (content/structure/language) and aggregates per-criterion scores via a chairman agent (average per criterion + veto '
    'flag when grader totals differ > 2.0), then applies rounding rules.'
)

# Experimental setup

doc.add_heading('4 Experimental Setup', level=1)
doc.add_paragraph(
    'Assignments: BT1–BT4 (rubric-based, total score 0–10). Two human graders score a subset; GT = (Ng1+Ng2)/2. '
    'Metrics: Inter-rater QWK, and AI vs GT QWK/MAE/Pearson.'
)

# Results

doc.add_heading('5 Results', level=1)

doc.add_heading('5.1 Human agreement and AI-vs-human metrics (human-scored subset)', level=2)

table = doc.add_table(rows=1, cols=9)
hdr = table.rows[0].cells
for i, h in enumerate(['BT', 'N_sample', 'Human QWK', 'P1 QWK', 'P1 MAE', 'P1 r', 'P2 QWK', 'P2 MAE', 'P2 r']):
    hdr[i].text = h

for _, r in summary.iterrows():
    row = table.add_row().cells
    row[0].text = str(r['BT'])
    row[1].text = str(int(r['N_human_sample']))
    row[2].text = f"{r['Inter_rater_QWK']:.3f}"
    row[3].text = f"{r['P1_vs_GT_QWK']:.3f}"
    row[4].text = f"{r['P1_vs_GT_MAE']:.3f}"
    row[5].text = f"{r['P1_vs_GT_Pearson_r']:.3f}"
    row[6].text = f"{r['P2_vs_GT_QWK']:.3f}"
    row[7].text = f"{r['P2_vs_GT_MAE']:.3f}"
    row[8].text = f"{r['P2_vs_GT_Pearson_r']:.3f}"

doc.add_paragraph('Observation: Phase 1 outperforms Phase 2 in QWK on all four assignments in the human-scored subset.')


doc.add_heading('5.2 Systematic score shift of Phase 2 vs Phase 1 (full dataset outputs)', level=2)

table2 = doc.add_table(rows=1, cols=5)
h2 = table2.rows[0].cells
for i, h in enumerate(['BT', 'N_full', 'mean(P2−P1)', 'median(P2−P1)', 'veto rate']):
    h2[i].text = h

for _, r in summary.iterrows():
    row = table2.add_row().cells
    row[0].text = str(r['BT'])
    row[1].text = str(int(r['N']))
    row[2].text = f"{r['mean_diff(P2-P1)']:.3f}"
    row[3].text = f"{r['median_diff(P2-P1)']:.3f}"
    row[4].text = f"{r['pct_veto']:.3f}"

# Discussion & recommendations

doc.add_heading('6 Discussion & Recommendations', level=1)
doc.add_paragraph(
    'We interpret multi-agent underperformance via (i) calibration mismatch and aggregation smoothing, (ii) assignment-dependent score shift, and '
    '(iii) surface-feature sensitivity vs rubric intent (especially for BT1–BT2 which emphasize factual completeness and sourcing).'
)

doc.add_paragraph('Recommendations before expecting gains from multi-agent grading:')
for rec in [
    'R1. Calibrate grader severity using anchor papers and learn per-grader offset/scale mapping before aggregation.',
    'R2. Use robust aggregation (median/trimmed mean) instead of simple per-criterion mean.',
    'R3. Rubric-aligned authority weighting: downweight language/style judgments for BT1–BT2; increase for BT3–BT4.',
]:
    doc.add_paragraph(rec)

# Limitations / TODO list based on critique

doc.add_heading('7 Submission Readiness Checklist (based on advisor/reviewer blockers)', level=1)
doc.add_paragraph('Blockers:')
doc.add_paragraph('- Write Related Work with real citations (15–20).')
doc.add_paragraph('- Add exact model configuration (model name/version, temperature, top_p, max tokens, prompt template).')
doc.add_paragraph('- Add properly formatted tables with N, QWK/MAE/Pearson and clear captions; optionally add confidence intervals.')
doc.add_paragraph('- Explain why BT1/BT2 are harder (assignment characteristics + rubric intent) using evidence (TC-level diffs; examples).')

doc.add_paragraph('Strengtheners:')
doc.add_paragraph('- Justify N=30 sampling (random/stratified); consider bootstrap CIs for QWK.')
doc.add_paragraph('- Add a brief description of each assignment and rubric criteria maxima.')
doc.add_paragraph('- Add statistical significance tests (bootstrap) or at least discuss limitations.')
doc.add_paragraph('- Add 1–2 figures: pipeline diagram; scatter plots (AI vs GT); score distribution plots.')


doc.add_heading('8 Conclusion', level=1)
doc.add_paragraph(
    'We provide an auditable OpenClaw grading pipeline and reliability-first evaluation on real course submissions, showing that multi-agent grading '
    'requires calibration and robust aggregation to improve reliability.'
)

# Save
out_path.parent.mkdir(parents=True, exist_ok=True)
doc.save(out_path)
print(f'Wrote: {out_path.resolve()}')
