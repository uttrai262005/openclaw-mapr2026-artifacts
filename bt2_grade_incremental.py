import json
import os
import re
from pathlib import Path
from typing import Any, Dict, List, Tuple

import urllib.request
from openpyxl import Workbook, load_workbook

WS = Path(__file__).resolve().parent
BT_DIR = WS / "dataset_clean" / "BT2"
RUBRIC_PATH = WS / "rubric" / "rubric_BT2.docx"
OUT_PATH = Path(os.environ.get("OUT_PATH", str(WS / "output" / "ket_qua_BT2.xlsx")))
OUT_PATH_RESUME = WS / "output" / "ket_qua_BT2_resume.xlsx"

MAX_TC1 = 1.5
MAX_TC2 = 2.0
MAX_TC3 = 2.5
MAX_TC4 = 2.0
MAX_TC5 = 2.0
MAX_TOTAL = 10.0

MSSV_RE = re.compile(r"(\d{6,})")


def round_quarter(x: float) -> float:
    return round(x * 4) / 4


def round_half(x: float) -> float:
    return round(x * 2) / 2


def clamp(x: float, lo: float, hi: float) -> float:
    return max(lo, min(hi, x))


def extract_mssv(path: Path) -> str:
    m = MSSV_RE.search(path.stem)
    return m.group(1) if m else path.stem


def read_docx_text(path: Path) -> str:
    from docx import Document

    d = Document(str(path))
    parts: List[str] = []
    for p in d.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for tb in d.tables:
        for row in tb.rows:
            cells = [(c.text or "").strip().replace("\n", " ") for c in row.cells]
            line = " | ".join([c for c in cells if c])
            if line:
                parts.append(line)
    return "\n".join(parts).strip()


def read_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(str(path))
        texts: List[str] = []
        for page in reader.pages:
            t = (page.extract_text() or "").strip()
            if t:
                texts.append(t)
        return "\n\n".join(texts).strip()
    except Exception:
        # pypdf failed; return empty to signal unreadable PDF in this pipeline.
        return ""


def read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore").strip()


def read_submission_text(path: Path) -> str:
    suf = path.suffix.lower()
    if suf == ".docx":
        return read_docx_text(path)
    if suf == ".pdf":
        return read_pdf_text(path)
    if suf in {".txt", ".md"}:
        return read_text_file(path)
    return read_text_file(path)


def get_gateway_config() -> Tuple[str, str, str]:
    url = os.environ.get("OPENCLAW_GATEWAY_URL", "").strip()
    token = os.environ.get("OPENCLAW_GATEWAY_TOKEN", "").strip()
    model = os.environ.get("OPENCLAW_MODEL", "").strip() or "openai-codex/gpt-5.2"

    cfg_path = Path(os.path.expanduser("~")) / ".openclaw" / "openclaw.json"
    cfg: Dict[str, Any] = {}
    if cfg_path.exists():
        try:
            cfg = json.loads(cfg_path.read_text(encoding="utf-8"))
        except Exception:
            cfg = {}

    def deep_get(d: Any, keys: List[str]) -> Any:
        cur = d
        for k in keys:
            if not isinstance(cur, dict) or k not in cur:
                return None
            cur = cur[k]
        return cur

    if not url:
        port = deep_get(cfg, ["gateway", "port"]) or deep_get(cfg, ["gateway", "http", "port"]) or 8317
        url = f"http://127.0.0.1:{int(port)}/v1/chat/completions"
    else:
        if url.endswith("/v1"):
            url = url + "/chat/completions"
        elif url.endswith("/v1/"):
            url = url[:-1] + "/chat/completions"
        elif url.endswith("/chat/completions"):
            pass
        else:
            url = url.rstrip("/") + "/v1/chat/completions"

    if not token:
        token = deep_get(cfg, ["gateway", "auth", "token"]) or deep_get(cfg, ["gateway", "token"]) or ""
        token = str(token) if token is not None else ""

    if not token:
        raise RuntimeError("Missing gateway token. Set OPENCLAW_GATEWAY_TOKEN or configure gateway.auth.token")

    return url, token, model


def call_llm_once(url: str, token: str, model: str, rubric_text: str, submission_text: str) -> Dict[str, Any]:
    # NOTE (English): The prompt is in Vietnamese because submissions/rubrics are Vietnamese.
    # Reproducibility relies on rubric-injection + JSON-only response + single call per submission.
    system = (
        "Bạn là giám khảo chấm bài BT2. Chỉ dùng đúng rubric được cung cấp. "
        "Chấm đủ 5 tiêu chí trong MỘT lần. "
        "Trả về DUY NHẤT một JSON hợp lệ (không markdown, không giải thích ngoài JSON)."
    )
    user = (
        "RUBRIC (trích từ rubric_BT2.docx):\n"
        f"{rubric_text}\n\n"
        "BÀI LÀM CỦA SINH VIÊN (text trích xuất):\n"
        f"{submission_text}\n\n"
        "YÊU CẦU OUTPUT JSON với schema:\n"
        "{\n"
        "  \"tc1\": <number 0..1.5>,\n"
        "  \"tc2\": <number 0..2>,\n"
        "  \"tc3\": <number 0..2.5>,\n"
        "  \"tc4\": <number 0..2>,\n"
        "  \"tc5\": <number 0..2>\n"
        "}\n"
        "Chỉ xuất JSON."
    )

    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "temperature": 0.1,
        "max_tokens": int(os.environ.get("OPENCLAW_MAX_TOKENS", "150")),
    }

    req = urllib.request.Request(
        url,
        data=json.dumps(payload).encode("utf-8"),
        headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {token}",
        },
        method="POST",
    )

    with urllib.request.urlopen(req, timeout=180) as resp:
        raw = resp.read().decode("utf-8", errors="ignore")

    data = json.loads(raw)
    content = data["choices"][0]["message"]["content"]
    if not isinstance(content, str):
        content = str(content)
    content = content.strip()

    try:
        return json.loads(content)
    except Exception:
        m = re.search(r"\{[\s\S]*\}", content)
        if not m:
            raise RuntimeError("Model did not return JSON")
        return json.loads(m.group(0))


def ensure_workbook(path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)
    if path.exists():
        return
    wb = Workbook()
    ws = wb.active
    ws.title = "BT2"
    ws.append(
        [
            "MSSV",
            "TC1(/1.5)",
            "TC2(/2)",
            "TC3(/2.5)",
            "TC4(/2)",
            "TC5(/2)",
            "Tổng(/10)",
            "Nhận xét ngắn",
            "Ghi chú",
        ]
    )
    wb.save(path)


def load_rubric_text() -> str:
    text = read_docx_text(RUBRIC_PATH)
    if not text:
        raise RuntimeError("Rubric extracted empty")
    return text[:12000]


def build_index(ws) -> Dict[str, int]:
    idx: Dict[str, int] = {}
    for r in range(2, ws.max_row + 1):
        v = ws.cell(r, 1).value
        if v is None:
            continue
        mssv = str(v).strip()
        if mssv and mssv not in idx:
            idx[mssv] = r
    return idx


def row_is_valid(ws, r: int) -> bool:
    total = ws.cell(r, 7).value
    note = ws.cell(r, 9).value
    if note is not None and str(note).strip().lower().startswith("lỗi"):
        return False
    if total is None:
        return False
    if str(total).strip() == "":
        return False
    return True


def main():
    if not BT_DIR.exists():
        raise SystemExit(f"BT2 dir not found: {BT_DIR}")
    if not RUBRIC_PATH.exists():
        raise SystemExit(f"Rubric not found: {RUBRIC_PATH}")

    ensure_workbook(OUT_PATH)
    out_path = OUT_PATH
    wb = load_workbook(out_path)
    ws = wb.active
    idx = build_index(ws)

    url, token, model = get_gateway_config()
    rubric_text = load_rubric_text()

    files = sorted([p for p in BT_DIR.rglob('*') if p.is_file() and p.suffix.lower() in {'.pdf','.docx','.txt','.md'}])
    print(f"BT2 files: {len(files)}")
    print(f"Writing incrementally to: {out_path}")

    skipped_valid = 0
    graded = 0
    recorded_errors = 0

    for i, path in enumerate(files, 1):
        mssv = extract_mssv(path)
        r = idx.get(mssv)
        if r is not None and row_is_valid(ws, r):
            skipped_valid += 1
            continue

        # Read submission text
        try:
            sub_text = read_submission_text(path)
            if not sub_text or len(sub_text.strip()) < 20:
                raise RuntimeError("Empty/too-short extracted text")
        except Exception:
            if r is None:
                ws.append([mssv, "", "", "", "", "", "", "", "Lỗi"])
                idx[mssv] = ws.max_row
            else:
                for c in range(2, 9):
                    ws.cell(r, c).value = ""
                ws.cell(r, 9).value = "Lỗi"
            try:
                wb.save(out_path)
            except PermissionError:
                out_path = OUT_PATH_RESUME
                wb.save(out_path)
            recorded_errors += 1
            print(f"[{i}/{len(files)}] {path.name}: READ ERROR -> recorded")
            continue

        # Exactly ONE API call per submission.
        try:
            res = call_llm_once(url, token, model, rubric_text, sub_text)
            tc1 = round_quarter(clamp(float(res.get("tc1", 0.0)), 0.0, MAX_TC1))
            tc2 = round_quarter(clamp(float(res.get("tc2", 0.0)), 0.0, MAX_TC2))
            tc3 = round_quarter(clamp(float(res.get("tc3", 0.0)), 0.0, MAX_TC3))
            tc4 = round_quarter(clamp(float(res.get("tc4", 0.0)), 0.0, MAX_TC4))
            tc5 = round_quarter(clamp(float(res.get("tc5", 0.0)), 0.0, MAX_TC5))
            total = round_half(clamp(tc1 + tc2 + tc3 + tc4 + tc5, 0.0, MAX_TOTAL))
            short = ""  # scores-only mode (no feedback field)

            if r is None:
                ws.append([mssv, tc1, tc2, tc3, tc4, tc5, total, short, ""])
                idx[mssv] = ws.max_row
            else:
                ws.cell(r, 2).value = tc1
                ws.cell(r, 3).value = tc2
                ws.cell(r, 4).value = tc3
                ws.cell(r, 5).value = tc4
                ws.cell(r, 6).value = tc5
                ws.cell(r, 7).value = total
                ws.cell(r, 8).value = short
                ws.cell(r, 9).value = ""

            try:
                wb.save(out_path)
            except PermissionError:
                out_path = OUT_PATH_RESUME
                wb.save(out_path)
            graded += 1
            print(f"[{i}/{len(files)}] {path.name}: OK total={total}")
        except Exception:
            if r is None:
                ws.append([mssv, "", "", "", "", "", "", "", "Lỗi"])
                idx[mssv] = ws.max_row
            else:
                for c in range(2, 9):
                    ws.cell(r, c).value = ""
                ws.cell(r, 9).value = "Lỗi"
            try:
                wb.save(out_path)
            except PermissionError:
                out_path = OUT_PATH_RESUME
                wb.save(out_path)
            recorded_errors += 1
            print(f"[{i}/{len(files)}] {path.name}: API/JSON ERROR -> recorded")

    print(f"SKIPPED valid: {skipped_valid}")
    print(f"GRADED this run: {graded}")
    print(f"ERROR recorded: {recorded_errors}")


if __name__ == "__main__":
    main()
