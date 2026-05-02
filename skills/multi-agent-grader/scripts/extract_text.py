from __future__ import annotations

import subprocess
from pathlib import Path
from typing import List, Optional


def extract_text(path: str | Path) -> str:
    path = Path(path)
    suf = path.suffix.lower()
    if suf in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")

    if suf == ".docx":
        from docx import Document  # type: ignore

        doc = Document(str(path))
        parts: List[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join((cell.text or "").strip() for cell in row.cells))
        return "\n".join(parts)

    if suf in {".png", ".jpg", ".jpeg", ".webp"}:
        return _ocr_image_best_effort(path)

    if suf == ".pdf":
        # Try fitz (PyMuPDF) first
        try:
            import fitz  # type: ignore

            doc = fitz.open(str(path))
            parts = []
            for page in doc:
                parts.append(page.get_text("text") or "")
            txt = "\n".join(parts)
            if txt.strip():
                return txt
        except Exception:
            pass

        # Try pdfplumber
        try:
            import pdfplumber  # type: ignore

            with pdfplumber.open(str(path)) as pdf:
                parts = [(p.extract_text() or "") for p in pdf.pages]
            txt = "\n".join(parts)
            if txt.strip():
                return txt
        except Exception:
            pass

        # Try pypdf
        try:
            from pypdf import PdfReader  # type: ignore

            reader = PdfReader(str(path))
            txt = "\n".join((page.extract_text() or "") for page in reader.pages)
            if txt.strip():
                return txt
        except Exception:
            pass

        # Fallback: node extractor if available
        js = Path(__file__).resolve().parents[2] / "_extract_pdf_text.js"
        if js.exists():
            try:
                out = subprocess.check_output(["node", str(js), str(path)], text=True, encoding="utf-8", errors="ignore")
                if out.strip():
                    return out
            except Exception:
                pass

        # Optional OCR (only if dependencies exist). Best-effort.
        txt = _ocr_pdf_best_effort(path)
        if txt.strip():
            return txt

        raise RuntimeError("Could not extract text from PDF (no extractor succeeded)")

    raise ValueError(f"Unsupported file type: {suf}")


def _ocr_image_best_effort(path: Path) -> str:
    """OCR for image files using pytesseract.

    Notes:
    - Requires Tesseract OCR installed.
    - We ship `vie.traineddata` in skills/giang-vien-v2/scripts/tessdata so OCR
      can work even if the system tesseract doesn't include Vietnamese.
    """
    try:
        import os
        import pytesseract  # type: ignore
        from PIL import Image  # type: ignore
    except Exception:
        return ""

    # Ensure tesseract.exe is discoverable even if PATH isn't refreshed
    tesseract_exe = Path("C:/Program Files/Tesseract-OCR/tesseract.exe")
    if tesseract_exe.exists():
        pytesseract.pytesseract.tesseract_cmd = str(tesseract_exe)

    bundled = Path(__file__).resolve().parent / "tessdata"
    has_vie = (bundled / "vie.traineddata").exists()
    has_eng = (bundled / "eng.traineddata").exists()

    try:
        img = Image.open(str(path)).convert("RGB")

        # light preprocessing to help OCR on screenshots
        try:
            from PIL import ImageOps  # type: ignore

            g = ImageOps.grayscale(img)
            g = ImageOps.autocontrast(g)
            img2 = g
        except Exception:
            img2 = img

        # Prefer bundled tessdata when available
        if bundled.exists() and (has_vie or has_eng):
            os.environ["TESSDATA_PREFIX"] = str(bundled.parent)

        config = "--psm 6"

        # Try Vietnamese+English if both available; else try Vietnamese only; else English
        try:
            if has_vie and has_eng:
                return pytesseract.image_to_string(img2, lang="vie+eng", config=config) or ""
            if has_vie:
                return pytesseract.image_to_string(img2, lang="vie", config=config) or ""
        except Exception:
            pass

        # Fallback: use system tessdata for English
        try:
            os.environ.pop("TESSDATA_PREFIX", None)
        except Exception:
            pass
        return pytesseract.image_to_string(img2, lang="eng", config=config) or ""

    except Exception:
        return ""


def _ocr_pdf_best_effort(path: Path) -> str:
    """Very best-effort OCR path: pdf -> images -> pytesseract.

    Will return "" if OCR is not available.
    """
    try:
        import fitz  # type: ignore
        from PIL import Image  # type: ignore
    except Exception:
        return ""

    try:
        doc = fitz.open(str(path))
        parts: List[str] = []
        for i, page in enumerate(doc):
            pix = page.get_pixmap(dpi=200)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            # route via image OCR (handles deps)
            try:
                import pytesseract  # type: ignore

                try:
                    parts.append(pytesseract.image_to_string(img, lang="vie+eng"))
                except Exception:
                    parts.append(pytesseract.image_to_string(img, lang="eng"))
            except Exception:
                parts.append("")
        return "\n".join(parts)
    except Exception:
        return ""
