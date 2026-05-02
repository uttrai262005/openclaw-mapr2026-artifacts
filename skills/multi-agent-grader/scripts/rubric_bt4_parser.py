from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Dict, Tuple

from schema import RubricBT4


def extract_text_from_docx(path: str | Path) -> str:
    path = Path(path)
    from docx import Document  # type: ignore

    d = Document(str(path))
    parts = []
    for p in d.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for tb in d.tables:
        for row in tb.rows:
            cells = [(c.text or "").strip().replace("\n", " ") for c in row.cells]
            line = " | ".join([c for c in cells if c])
            if line:
                parts.append(line)
    return "\n".join(parts)


_TC_ROW_RE = re.compile(
    r"^\s*(?P<tc>[1-4])\.(?:\s*)[^|]*\|\s*(?P<max>\d+(?:[\.,]\d+)?)\s*\|",
    re.IGNORECASE,
)


def parse_rubric_bt4(path: str | Path) -> Tuple[RubricBT4, Dict[str, Any]]:
    """Parse rubric_BT4.docx to RubricBT4 + raw metadata.

    This is intentionally minimal and robust: it only needs the max points per TC.
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(path)
    if path.suffix.lower() != ".docx":
        raise ValueError("BT4 rubric parser expects .docx")

    text = extract_text_from_docx(path)
    maxima = {"1": None, "2": None, "3": None, "4": None}
    for line in text.splitlines():
        m = _TC_ROW_RE.match(line.strip())
        if not m:
            continue
        tc = m.group("tc")
        mx = float(m.group("max").replace(",", "."))
        maxima[tc] = mx

    rubric = RubricBT4(
        tc1_max=float(maxima["1"] or 3.0),
        tc2_max=float(maxima["2"] or 4.0),
        tc3_max=float(maxima["3"] or 2.5),
        tc4_max=float(maxima["4"] or 0.5),
    )
    raw = {
        "source": str(path),
        "extracted_preview": "\n".join(text.splitlines()[:20]),
        "maxima": maxima,
    }
    return rubric, raw
